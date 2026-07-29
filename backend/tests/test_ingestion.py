import io
import pytest
import docx
from reportlab.pdfgen import canvas
from app.services.extractor import extract_from_pdf, extract_from_docx, extract_from_txt

@pytest.fixture
def sample_pdf():
    # Generate a dummy PDF in memory
    packet = io.BytesIO()
    can = canvas.Canvas(packet)
    can.drawString(10, 800, "ARTICLE 1")
    can.drawString(10, 780, "This is the first paragraph of the legal document.")
    can.showPage()
    can.drawString(10, 800, "This is page 2.")
    can.save()
    packet.seek(0)
    return packet.read()

@pytest.fixture
def sample_docx():
    # Generate a dummy DOCX in memory
    doc = docx.Document()
    doc.add_heading('Section 1', level=1)
    doc.add_paragraph('This is a test paragraph in docx.')
    
    packet = io.BytesIO()
    doc.save(packet)
    packet.seek(0)
    return packet.read()

@pytest.fixture
def sample_txt():
    return b"Heading 1\n\nThis is some text in txt format."

def test_ingest_pdf(sample_pdf):
    sections = extract_from_pdf(sample_pdf)
    assert len(sections) > 0
    first_section = sections[0]
    
    # reportlab might not put a double newline between the text elements
    assert "ARTICLE 1" in first_section["text"] or first_section["heading"] == "ARTICLE 1"
    assert "first paragraph" in first_section["text"]
    assert first_section["page"] == 1

def test_ingest_docx(sample_docx):
    sections = extract_from_docx(sample_docx)
    assert len(sections) > 0
    first_section = sections[0]
    assert first_section["heading"] == "Section 1"
    assert "test paragraph" in first_section["text"]

def test_ingest_txt(sample_txt):
    sections = extract_from_txt(sample_txt)
    assert len(sections) > 0
    first_section = sections[0]
    assert "Heading 1" in first_section["heading"]
    assert "text in txt format" in first_section["text"]
