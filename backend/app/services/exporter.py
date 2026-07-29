"""Export structured legal reports to PDF and DOCX."""

from __future__ import annotations

import io
from typing import Any, Dict

from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


def _safe(text: Any) -> str:
    return str(text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def build_pdf(report: Dict[str, Any], filename: str) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph(_safe(f"ClauseGuard Report — {filename}"), styles["Title"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph(_safe(report.get("disclaimer") or ""), styles["Normal"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Executive Summary", styles["Heading2"]))
    story.append(Paragraph(_safe(report.get("executive_summary") or "N/A"), styles["Normal"]))
    story.append(Spacer(1, 8))
    story.append(Paragraph(_safe(f"Overall risk: {report.get('overall_risk', 'unknown')}"), styles["Normal"]))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Review Priorities", styles["Heading2"]))
    for item in report.get("review_priorities") or []:
        story.append(
            Paragraph(
                _safe(f"• {item.get('title')}: {item.get('rationale')} Action: {item.get('action')}"),
                styles["Normal"],
            )
        )
    story.append(Spacer(1, 12))

    story.append(Paragraph("Obligations", styles["Heading2"]))
    for item in report.get("obligations") or []:
        story.append(
            Paragraph(
                _safe(
                    f"• [{item.get('status', 'unconfirmed')}] {item.get('party') or 'party'}: "
                    f"{item.get('action')} (deadline={item.get('deadline') or item.get('period') or 'n/a'})"
                ),
                styles["Normal"],
            )
        )
    story.append(Spacer(1, 12))

    story.append(Paragraph("Negotiation Playbook", styles["Heading2"]))
    for item in report.get("negotiation_playbook") or []:
        story.append(
            Paragraph(
                _safe(
                    f"• {item.get('category')}: Ask={item.get('primary_ask')} "
                    f"Fallback={item.get('fallback')}"
                ),
                styles["Normal"],
            )
        )
    story.append(Spacer(1, 12))

    story.append(Paragraph("Flagged Clauses", styles["Heading2"]))
    for flag in report.get("flags") or []:
        story.append(
            Paragraph(
                _safe(
                    f"• [{flag.get('severity')}] {flag.get('category')} (p.{flag.get('page')}): "
                    f"{flag.get('explanation')}"
                ),
                styles["Normal"],
            )
        )
        story.append(Paragraph(_safe(flag.get("text") or ""), styles["Italic"]))
        story.append(Spacer(1, 6))

    story.append(Spacer(1, 16))
    story.append(
        Paragraph(
            _safe(f"Analyzed at {report.get('analyzed_at') or 'n/a'} using {report.get('model') or 'n/a'}"),
            styles["Normal"],
        )
    )

    doc.build(story)
    return buffer.getvalue()


def build_docx(report: Dict[str, Any], filename: str) -> bytes:
    document = DocxDocument()
    document.add_heading(f"ClauseGuard Report — {filename}", level=1)
    document.add_paragraph(report.get("disclaimer") or "")
    document.add_heading("Executive Summary", level=2)
    document.add_paragraph(report.get("executive_summary") or "N/A")
    document.add_paragraph(f"Overall risk: {report.get('overall_risk', 'unknown')}")

    document.add_heading("Review Priorities", level=2)
    for item in report.get("review_priorities") or []:
        document.add_paragraph(
            f"{item.get('title')}: {item.get('rationale')} Action: {item.get('action')}",
            style="List Bullet",
        )

    document.add_heading("Obligations", level=2)
    for item in report.get("obligations") or []:
        document.add_paragraph(
            f"[{item.get('status', 'unconfirmed')}] {item.get('party') or 'party'}: "
            f"{item.get('action')} (deadline={item.get('deadline') or item.get('period') or 'n/a'})",
            style="List Bullet",
        )

    document.add_heading("Negotiation Playbook", level=2)
    for item in report.get("negotiation_playbook") or []:
        document.add_paragraph(
            f"{item.get('category')}: Ask={item.get('primary_ask')} | Fallback={item.get('fallback')}",
            style="List Bullet",
        )

    document.add_heading("Flagged Clauses", level=2)
    for flag in report.get("flags") or []:
        document.add_paragraph(
            f"[{flag.get('severity')}] {flag.get('category')} (p.{flag.get('page')}): {flag.get('explanation')}",
            style="List Bullet",
        )
        document.add_paragraph(flag.get("text") or "")

    document.add_paragraph(
        f"Analyzed at {report.get('analyzed_at') or 'n/a'} using {report.get('model') or 'n/a'}"
    )

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
