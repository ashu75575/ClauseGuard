"""
Document Extractor – Extracts text from PDF, DOCX, and TXT files.

Improvements:
  - Expanded boilerplate filter (signature blocks, page numbers, headers)
  - Better heading detection (ALL-CAPS, colon-terminated, known legal titles)
  - Table extraction for PDFs (fee schedules, SLA matrices)
  - Encoding detection for TXT files (handles non-UTF-8)
  - DOCX page estimation via rendered page breaks
  - Structured logging instead of print()
"""

import io
import uuid
import re
import logging
from typing import List, Optional

import pdfplumber
import docx
import pytesseract

logger = logging.getLogger(__name__)

# --- Boilerplate patterns (expanded) ---
BOILERPLATE_PATTERNS = [
    r"^last updated",
    r"^\S+@\S+\.\S+$",  # bare email address
    r"^if you have questions",
    r"^page\s+\d+",  # page numbers
    r"^confidential",
    r"^_+$",  # underscores (signature lines)
    r"^signature:?\s*$",
    r"^by:\s*.*date:\s*",
    r"^\d{1,2}[/-]\d{1,2}[/-]\d{2,4}$",  # bare dates
    r"^all rights reserved",
    r"^copyright\s+",
    r"^table of contents",
    r"^\[?\s*initials?\s*\]?\s*$",
    r"^witness:",
]

# Known legal section titles that should be treated as headings
KNOWN_LEGAL_HEADINGS = {
    "definitions",
    "term and termination",
    "limitation of liability",
    "indemnification",
    "confidentiality",
    "governing law",
    "dispute resolution",
    "arbitration",
    "intellectual property",
    "data protection",
    "privacy policy",
    "force majeure",
    "assignment",
    "entire agreement",
    "severability",
    "amendments",
    "notices",
    "representations and warranties",
    "payment terms",
    "fees",
    "non-compete",
    "non-solicitation",
    "termination",
    "renewal",
    "warranty",
    "disclaimer",
    "general provisions",
    "miscellaneous",
    "scope of services",
    "obligations",
    "insurance",
}


def is_boilerplate(text: str) -> bool:
    """
    Returns True for front-matter/closing lines that aren't real clauses.
    """
    text_lower = text.strip().lower()

    # Very short lines (1-2 words) that aren't headings
    words = text_lower.split()
    if len(words) <= 1 and not text_lower.replace(".", "").replace(":", "").strip():
        return True

    return any(re.match(p, text_lower) for p in BOILERPLATE_PATTERNS)


def is_heading(text: str) -> bool:
    """
    Heuristic to determine if a text block is a section heading.

    Detects:
      - Numbered clauses: "1.", "1.1", "1.1.2 Term"
      - Article/Section markers: "Article 5", "Section 3.1"
      - ALL-CAPS titles under 80 chars: "LIMITATION OF LIABILITY"
      - Colon-terminated short lines: "Governing Law:"
      - Known legal section titles
    """
    text = text.strip()
    if not text:
        return False
    if "\n" in text:
        return False
    if len(text) > 100:
        return False

    # Numbered clause: "1.", "1.1", "1.1.2 Term", "3.4.1. Obligations"
    if re.match(r"^\d+(\.\d+)*\.?\s+\S", text):
        return True

    # Article/Section/Heading markers
    if re.match(r"^(heading|article|section|clause|schedule|exhibit|appendix|annex)\s+\d+", text, re.IGNORECASE):
        return True

    # ALL-CAPS title (at least 2 words, under 80 chars, mostly uppercase)
    if len(text) <= 80 and len(text.split()) >= 2:
        alpha_chars = [c for c in text if c.isalpha()]
        if alpha_chars and sum(1 for c in alpha_chars if c.isupper()) / len(alpha_chars) > 0.8:
            return True

    # Colon-terminated short title (under 50 chars)
    if text.endswith(":") and len(text) <= 50 and len(text.split()) <= 6:
        return True

    # Known legal heading (exact match, case-insensitive)
    text_normalised = re.sub(r"[\d.:\-–—]+", "", text).strip().lower()
    if text_normalised in KNOWN_LEGAL_HEADINGS:
        return True

    return False


def _serialize_table(table: list) -> str:
    """Convert a pdfplumber table to a readable text representation."""
    if not table:
        return ""
    rows = []
    for row in table:
        cells = [str(cell).strip() if cell else "" for cell in row]
        rows.append(" | ".join(cells))
    return "\n".join(rows)


def extract_from_pdf(file_bytes: bytes) -> List[dict]:
    """
    Extracts text blocks, headings, and tables from a PDF byte stream.
    Falls back to OCR for scanned pages.
    """
    sections = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for i, page in enumerate(pdf.pages):
            page_num = i + 1

            # --- Extract tables first ---
            try:
                tables = page.extract_tables()
                for table in tables:
                    table_text = _serialize_table(table)
                    if table_text and len(table_text.strip()) > 20:
                        sections.append(
                            {
                                "section_id": str(uuid.uuid4()),
                                "heading": f"Table (Page {page_num})",
                                "text": table_text,
                                "page": page_num,
                            }
                        )
            except Exception as e:
                logger.warning("Table extraction failed for page %d: %s", page_num, e)

            # --- Extract text ---
            text = page.extract_text()

            if not text or len(text.strip()) < 50:
                try:
                    img = page.to_image(resolution=300).original
                    text = pytesseract.image_to_string(img)
                    if text and len(text.strip()) >= 50:
                        logger.info("OCR succeeded for page %d (%d chars)", page_num, len(text))
                except Exception as e:
                    logger.warning("OCR failed for page %d: %s", page_num, e)

            if text:
                blocks = re.split(r"\n\s*\n", text)
                current_heading = None

                for block in blocks:
                    block = block.strip()
                    if not block or is_boilerplate(block):
                        continue

                    if is_heading(block):
                        current_heading = block
                    else:
                        sections.append(
                            {
                                "section_id": str(uuid.uuid4()),
                                "heading": current_heading,
                                "text": block,
                                "page": page_num,
                            }
                        )

    logger.info("PDF extraction: %d pages → %d sections", len(pdf.pages) if hasattr(pdf, 'pages') else 0, len(sections))
    return sections


def extract_from_docx(file_bytes: bytes) -> List[dict]:
    """
    Extracts text blocks and headings from a DOCX byte stream.
    Estimates page numbers using rendered page break markers.
    """
    sections = []
    doc = docx.Document(io.BytesIO(file_bytes))
    current_heading = None
    estimated_page = 1
    char_count = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or is_boilerplate(text):
            continue

        # Estimate page breaks: check XML for lastRenderedPageBreak
        try:
            para_xml = para._element.xml
            if "lastRenderedPageBreak" in para_xml or "w:pageBreakBefore" in para_xml:
                estimated_page += 1
        except Exception:
            pass

        # Also estimate by character count (~3000 chars per page)
        char_count += len(text)
        page_by_chars = max(1, (char_count // 3000) + 1)
        page = max(estimated_page, page_by_chars)

        if para.style.name.startswith("Heading") or is_heading(text):
            current_heading = text
        else:
            sections.append(
                {
                    "section_id": str(uuid.uuid4()),
                    "heading": current_heading,
                    "text": text,
                    "page": page,
                }
            )

    logger.info("DOCX extraction: %d paragraphs → %d sections", len(doc.paragraphs), len(sections))
    return sections


def extract_from_txt(file_bytes: bytes) -> List[dict]:
    """
    Extracts text blocks and headings from a TXT byte stream.
    Detects encoding automatically and estimates page numbers.
    """
    # Try UTF-8 first, fall back to latin-1 (which never fails)
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = file_bytes.decode("latin-1")
            logger.info("TXT file decoded with latin-1 fallback")
        except Exception:
            text = file_bytes.decode("utf-8", errors="replace")
            logger.warning("TXT file decoded with replacement characters")

    sections = []
    blocks = re.split(r"\n\s*\n", text)
    current_heading = None
    char_count = 0

    for block in blocks:
        block = block.strip()
        if not block or is_boilerplate(block):
            continue

        char_count += len(block)
        estimated_page = max(1, (char_count // 3000) + 1)

        if is_heading(block):
            current_heading = block
        else:
            sections.append(
                {
                    "section_id": str(uuid.uuid4()),
                    "heading": current_heading,
                    "text": block,
                    "page": estimated_page,
                }
            )

    logger.info("TXT extraction: %d blocks → %d sections", len(blocks), len(sections))
    return sections
